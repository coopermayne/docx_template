"""
Script to create a sample Word document template with jinja2 placeholders
that can be used with docxtpl.
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_sample_template():
    """Create a sample .docx template with placeholders"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Document Template', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add a paragraph with information
    doc.add_paragraph('This is a sample document generated from a template.')
    doc.add_paragraph()
    
    # Add template variables
    doc.add_heading('Template Variables:', level=1)
    
    # Create a table with sample placeholders
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Value'
    
    # Add rows with jinja2 template variables
    fields = [
        ('Name', '{{ name }}'),
        ('Company', '{{ company }}'),
        ('Date', '{{ date }}'),
        ('Description', '{{ description }}'),
        ('Generated Date', '{{ generated_date }}')
    ]
    
    for field, value in fields:
        row_cells = table.add_row().cells
        row_cells[0].text = field
        row_cells[1].text = value
    
    doc.add_paragraph()
    
    # Add additional content section
    doc.add_heading('Additional Information:', level=1)
    doc.add_paragraph('{{ additional_info }}')
    
    doc.add_paragraph()
    doc.add_paragraph('---')
    doc.add_paragraph('Note: This document was automatically generated.')
    
    # Save the template
    template_dir = os.path.join(os.path.dirname(__file__), 'templates')
    os.makedirs(template_dir, exist_ok=True)
    template_path = os.path.join(template_dir, 'default_template.docx')
    doc.save(template_path)
    print(f"Sample template created at: {template_path}")

if __name__ == '__main__':
    create_sample_template()
