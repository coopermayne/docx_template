import os
import tempfile
from datetime import datetime
from typing import List, Dict, Any, Optional
from docxtpl import DocxTemplate
from models import Session, RFPRequest, Document
from config import Config
from api.objections import load_preset


class DocumentGenerator:
    """Generate Word documents for RFP responses."""

    def __init__(self):
        self.template_dir = Config.WORD_TEMPLATE_FOLDER
        self._uploaded_template_path = None  # Track for cleanup

    def generate_response(
        self,
        session: Session,
        court_name: str = "Superior Court of California",
        header_plaintiffs: str = "PLAINTIFF",
        header_defendants: str = "DEFENDANT",
        case_no: str = "",
        client_name: str = "Plaintiff",
        requesting_party: str = "Defendant",
        propounding_party: str = "Defendant",
        responding_party: str = "Plaintiff",
        set_number: str = "ONE",
        document_title: str = "",
        multiple_plaintiffs: bool = False,
        multiple_defendants: bool = False,
        multiple_propounding_parties: bool = False,
        multiple_responding_parties: bool = False,
        include_reasoning: bool = False
    ) -> str:
        """
        Generate an RFP response document.

        Args:
            session: The session containing requests and documents
            court_name: Name of the court
            header_plaintiffs: Plaintiffs for the case caption
            header_defendants: Defendants for the case caption
            case_no: Case number
            client_name: Name of the client/plaintiff
            requesting_party: Name of the party requesting production
            propounding_party: Name of the propounding party (starts with Defendant/Defendants)
            responding_party: Name of the responding party (starts with Plaintiff/Plaintiffs)
            set_number: The set number (e.g., "ONE", "TWO")
            document_title: The formal document title for the response
            multiple_plaintiffs: True if multiple plaintiffs in case caption
            multiple_defendants: True if multiple defendants in case caption
            multiple_propounding_parties: True if RFP propounded by multiple defendants
            multiple_responding_parties: True if RFP addressed to multiple plaintiffs
            include_reasoning: If True, include AI-generated arguments after each objection

        Returns:
            Path to the generated document
        """
        # Load objections preset
        preset = load_preset(session.objection_preset_id or 'default')
        objections_map = {}
        if preset:
            for obj in preset.get('objections', []):
                objections_map[obj['id']] = obj

        # Build documents map
        documents_map = {doc.id: doc for doc in session.documents}

        # Build context for template
        # Create short versions for document properties (255 char limit)
        def truncate_name(name: str, max_len: int = 50) -> str:
            """Truncate a name, trying to keep it meaningful."""
            if len(name) <= max_len:
                return name
            # Try to cut at a sensible point (semicolon, comma, or space)
            for sep in [';', ',', ' ']:
                if sep in name[:max_len]:
                    idx = name[:max_len].rfind(sep)
                    if idx > 10:  # Don't truncate too short
                        return name[:idx].strip() + ', et al.'
            return name[:max_len-3] + '...'

        # Generate default document title if not provided
        if not document_title:
            document_title = f"{responding_party.upper()}'S RESPONSES TO {propounding_party.upper()}'S {set_number} SET OF REQUESTS FOR PRODUCTION OF DOCUMENTS"

        context = {
            'court_name': court_name,
            'header_plaintiffs': header_plaintiffs,
            'header_defendants': header_defendants,
            'case_no': case_no,
            'client_name': client_name,
            'client_name_short': truncate_name(client_name),
            'requesting_party': requesting_party,
            'requesting_party_short': truncate_name(requesting_party),
            'propounding_party': propounding_party,
            'propounding_party_or_parties': "Defendants" if multiple_propounding_parties else "Defendant",
            'responding_party': responding_party,
            'set_number': set_number,
            'document_title': document_title,
            'multiple_plaintiffs': multiple_plaintiffs,
            'multiple_defendants': multiple_defendants,
            'multiple_propounding_parties': multiple_propounding_parties,
            'multiple_responding_parties': multiple_responding_parties,
            'date': datetime.now().strftime('%B %d, %Y'),
            'requests': []
        }

        # Process each request
        for req in session.requests:
            if not req.include_in_response:
                continue

            # Gather selected objections with full data
            selected_objections = []
            for obj_id in req.selected_objections:
                obj = objections_map.get(obj_id)
                if obj:
                    selected_objections.append(obj)

            # Gather selected documents with full data
            selected_documents = []
            for doc_id in req.selected_documents:
                doc = documents_map.get(doc_id)
                if doc:
                    selected_documents.append({
                        'id': doc.id,
                        'filename': doc.filename,
                        'bates_start': doc.bates_start,
                        'bates_end': doc.bates_end,
                        'description': doc.description
                    })

            # Build the response text from objections and documents
            response_text = self._build_response_text(
                selected_objections,
                selected_documents,
                responding_party,
                objection_arguments=req.objection_arguments if include_reasoning else None
            )

            # Build request data for template
            request_data = {
                'number': req.number,
                'question': req.text,
                'response': response_text,
                # Keep these for backwards compatibility
                'text': req.text,
                'objections': [{'id': obj['id'], 'name': obj['name'], 'formal_language': obj['formal_language']} for obj in selected_objections],
                'documents': selected_documents
            }

            context['requests'].append(request_data)

        # Try to get uploaded template from Supabase first
        template_path = None
        try:
            from api.templates import get_latest_template_path
            self._uploaded_template_path = get_latest_template_path('rfp')
            if self._uploaded_template_path:
                template_path = self._uploaded_template_path
        except Exception:
            pass

        # Fall back to local templates
        if not template_path:
            template_path = os.path.join(self.template_dir, 'rfp_template.docx')
            if not os.path.exists(template_path):
                template_path = os.path.join(self.template_dir, 'rfp_response.docx')

        if template_path and os.path.exists(template_path):
            doc = DocxTemplate(template_path)
            doc.render(context)
        else:
            # Generate document programmatically if no template
            return self._generate_without_template(context)

        # Save to temp file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()

        # Clean up uploaded template if used
        if self._uploaded_template_path:
            try:
                os.unlink(self._uploaded_template_path)
            except (OSError, FileNotFoundError):
                pass
            self._uploaded_template_path = None

        return temp_file.name

    def _build_response_text(
        self,
        objections: List[Dict],
        documents: List[Dict],
        responding_party: str,
        objection_arguments: Optional[Dict[str, str]] = None
    ) -> str:
        """
        Build the response text from objections and documents.

        Args:
            objections: List of selected objection dictionaries
            documents: List of selected document dictionaries
            responding_party: Name of the responding party
            objection_arguments: Optional dict of objection_id -> persuasive argument text

        Returns:
            Formatted response text
        """
        parts = []

        # Add objections (keep "Responding Party" as-is in formal language)
        if objections:
            objection_texts = []
            for obj in objections:
                text = obj['formal_language']
                # Append persuasive argument if available
                if objection_arguments and obj['id'] in objection_arguments:
                    argument = objection_arguments[obj['id']]
                    if argument:
                        text = f"{text} {argument}"
                objection_texts.append(text)
            parts.append(" ".join(objection_texts))

        # Add document production statement
        if documents:
            doc_parts = []
            for doc in documents:
                bates_str = ""
                if doc.get('bates_start'):
                    bates_str = f" ({doc['bates_start']}"
                    if doc.get('bates_end'):
                        bates_str += f"-{doc['bates_end']}"
                    bates_str += ")"
                doc_parts.append(f"{doc['filename']}{bates_str}")

            # Join documents with commas and "and" for last item
            if len(doc_parts) == 1:
                docs_text = doc_parts[0]
            elif len(doc_parts) == 2:
                docs_text = f"{doc_parts[0]} and {doc_parts[1]}"
            else:
                docs_text = ", ".join(doc_parts[:-1]) + f", and {doc_parts[-1]}"

            if objections:
                parts.append(f"Subject to and without waiving the foregoing objections, {responding_party} will produce the following documents responsive to this Request: {docs_text}.")
            else:
                parts.append(f"{responding_party} will produce the following documents responsive to this Request: {docs_text}.")

        # If no objections and no documents
        if not objections and not documents:
            parts.append(f"{responding_party} responds that there are no documents responsive to this Request.")

        return " ".join(parts)

    def _generate_without_template(self, context: Dict[str, Any]) -> str:
        """Generate document programmatically without a template."""
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Set up styles
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)

        # Title
        title = doc.add_heading('RESPONSES TO REQUESTS FOR PRODUCTION OF DOCUMENTS', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Header info
        doc.add_paragraph(f"PROPOUNDING PARTY: {context['propounding_party']}")
        doc.add_paragraph(f"RESPONDING PARTY: {context['responding_party']}")
        doc.add_paragraph(f"SET NUMBER: {context['set_number']}")
        doc.add_paragraph()

        # Each request
        for req in context['requests']:
            # Request header
            doc.add_heading(f"REQUEST NO. {req['number']}:", level=2)
            doc.add_paragraph(req['text'])
            doc.add_paragraph()

            # Response header
            doc.add_heading(f"RESPONSE TO REQUEST NO. {req['number']}:", level=2)

            # Objections - all in one paragraph
            if req['objections']:
                # Combine all objection formal language into one paragraph
                objection_texts = [obj['formal_language'] for obj in req['objections']]
                combined_objections = " ".join(objection_texts)
                doc.add_paragraph(combined_objections)
                doc.add_paragraph()

            # Responsive documents - in one sentence
            if req['documents']:
                # Build document list as inline text
                doc_parts = []
                for doc_ref in req['documents']:
                    bates_str = ""
                    if doc_ref.get('bates_start'):
                        bates_str = f" ({doc_ref['bates_start']}"
                        if doc_ref.get('bates_end'):
                            bates_str += f"-{doc_ref['bates_end']}"
                        bates_str += ")"
                    doc_parts.append(f"{doc_ref['filename']}{bates_str}")

                # Join documents with commas and "and" for last item
                if len(doc_parts) == 1:
                    docs_text = doc_parts[0]
                elif len(doc_parts) == 2:
                    docs_text = f"{doc_parts[0]} and {doc_parts[1]}"
                else:
                    docs_text = ", ".join(doc_parts[:-1]) + f", and {doc_parts[-1]}"

                if req['objections']:
                    doc.add_paragraph(f"Subject to and without waiving the foregoing objections, {context['responding_party']} will produce the following documents responsive to this Request: {docs_text}.")
                else:
                    doc.add_paragraph(f"{context['responding_party']} will produce the following documents responsive to this Request: {docs_text}.")
                doc.add_paragraph()

            # No objections and no documents
            if not req['objections'] and not req['documents']:
                doc.add_paragraph(f"{context['responding_party']} responds that there are no documents responsive to this Request.")

            doc.add_paragraph()

        # Signature block
        doc.add_paragraph()
        doc.add_paragraph(f"DATED: {context['date']}")
        doc.add_paragraph()
        doc.add_paragraph(context['responding_party'])

        # Save to temp file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()

        return temp_file.name


# Global instance
document_generator = DocumentGenerator()
